from docx import Document
import os
import glob as gb
# d = Document('./demo.docx')
def replace_text(old_text, new_text,path):
    for p in d.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in inline:
                if old_text in i.text:
                    text = i.text.replace(old_text, new_text)
                    i.text = text
                    i.add_picture(path)

img_path = gb.glob(".\\*.jpg")
d = Document('./demo.docx')
try:
    for path in img_path:
        print(path)
        dict_og=path.replace(".\\","")
        print(dict_og)
        
        replace_text(dict_og,"替换",path=dict_og)

except Exception as e:
    print(e)

finally:
    d.save('./demo2.docx')
